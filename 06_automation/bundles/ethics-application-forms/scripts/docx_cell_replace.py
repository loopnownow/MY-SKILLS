# -*- coding: utf-8 -*-
"""Preserve-format text replace in a .docx (paragraphs + unique table cells).

Usage:
  python docx_cell_replace.py <file.docx> --map map.json
  python docx_cell_replace.py <file.docx> <old> <new>
"""
from __future__ import annotations

import json
import sys


def set_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.text = text


def apply_map_to_text(text: str, mapping: dict[str, str]) -> str:
    out = text
    # longer keys first
    for old in sorted(mapping.keys(), key=len, reverse=True):
        if old in out:
            out = out.replace(old, mapping[old])
    return out


def unique_cells(table):
    seen = set()
    for row in table.rows:
        for cell in row.cells:
            tid = id(cell._tc)
            if tid in seen:
                continue
            seen.add(tid)
            yield cell


def process(path: str, mapping: dict[str, str]) -> None:
    from docx import Document

    doc = Document(path)
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs) if p.runs else p.text
        new = apply_map_to_text(full, mapping)
        if new != full:
            set_para_text(p, new)
    for table in doc.tables:
        for cell in unique_cells(table):
            full = cell.text
            new = apply_map_to_text(full, mapping)
            if new != full:
                paras = cell.paragraphs
                if not paras:
                    cell.text = new
                    continue
                set_para_text(paras[0], new)
                for p in paras[1:]:
                    set_para_text(p, "")
    doc.save(path)
    print("saved", path)


def main(argv: list[str]) -> int:
    if len(argv) < 3:
        print(__doc__)
        return 2
    path = argv[1]
    if argv[2] == "--map":
        with open(argv[3], encoding="utf-8") as f:
            mapping = json.load(f)
    else:
        mapping = {argv[2]: argv[3]}
    process(path, mapping)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
